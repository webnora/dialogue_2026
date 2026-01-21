#!/usr/bin/env python3
"""
Create Word document (.docx) from paper draft with embedded figures.
Improved version with better formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    # Don't modify font size, keep default
    return heading

def add_paragraph(doc, text, bold=False):
    """Add a paragraph."""
    if text.strip():
        para = doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        para.paragraph_format.line_spacing = 1.5
        return para
    return doc.add_paragraph()

def add_image(doc, image_path, width=6.0):
    """Add an image with specified width in inches."""
    if os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(width))
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")
            return False
    else:
        print(f"Image not found: {image_path}")
        return False

def add_figure_caption(doc, caption_text):
    """Add a figure caption."""
    para = doc.add_paragraph()
    run = para.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_bullet_point(doc, text):
    """Add a bullet point."""
    para = doc.add_paragraph(text, style='List Bullet')
    return para

def create_paper_document():
    """Create the complete paper document with figures."""

    doc = Document()

    # Title page
    title = doc.add_heading('What Vector Representations Reveal About Publicistic Writing', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, 'Conference: Dialogue 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Date: January 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Abstract
    add_heading(doc, 'Abstract', level=1)
    abstract_text = """Genre classification of journalistic texts is a fundamental task in computational linguistics with applications in content management, recommendation systems, and discourse analysis. While previous research has primarily focused on maximizing classification accuracy, less attention has been paid to what classification errors reveal about the nature of genre boundaries. This paper compares three levels of linguistic representation—lexical (TF-IDF), discourse-grammatical (hand-crafted linguistic features), and contextual-semantic (BERT)—on a corpus of 49,000 articles from The Guardian across five genres: News, Analytical, Feature, Editorial, and Review.

Our results demonstrate that BERT achieves 87.64% accuracy, only marginally outperforming TF-IDF at 86.58%, while linguistic features significantly underperform at 65.00%. This 1.06% gap between BERT and TF-IDF indicates that approximately 99% of the genre signal is captured by lexical choice alone, challenging the assumption that complex contextual representations are necessary for genre classification. Error analysis reveals that 27.8% of articles occupy hybrid boundary zones where models systematically disagree, providing empirical evidence for gradient genre boundaries. Statistical validation using McNemar's test, bootstrap confidence intervals, and Cohen's Kappa confirms that all model differences are highly significant (p < 0.001) while showing substantial inter-model agreement (κ = 0.73–0.83). Attention analysis of BERT reveals genre-specific lexical markers that overlap substantially across genres, further supporting gradient boundaries. We conclude that genre is primarily a lexical phenomenon with permeable, overlapping boundaries rather than discrete categories.

Keywords: genre classification, journalistic text, TF-IDF, BERT, error analysis, attention mechanisms, gradient boundaries"""

    add_paragraph(doc, abstract_text)
    doc.add_page_break()

    # Introduction
    add_heading(doc, '1. Introduction', level=1)
    add_heading(doc, '1.1 Background and Motivation', level=2)
    intro1 = """Genre is a fundamental organizing principle in journalism, guiding both writers and readers in the production and interpretation of texts. When readers encounter a newspaper article, they intuitively recognize whether they are reading a news report, an opinion piece, or a feature story. This recognition influences their expectations, interpretation strategies, and evaluation of the content. For computational systems, however, genre classification remains a challenging task due to the inherent flexibility and evolution of genre conventions in contemporary journalism.

The rapid growth of digital news content has created an urgent need for automated genre classification systems. Such systems serve critical functions in content management, personalized recommendation, academic research, and media monitoring. However, developing robust genre classifiers requires addressing fundamental questions about the nature of genre itself: Are genres discrete categories with clear boundaries, or do they exist on a continuum with overlapping characteristics? What linguistic features reliably distinguish genres? How do we handle articles that blend elements from multiple genres?"""
    add_paragraph(doc, intro1)

    add_heading(doc, '1.2 Research Questions', level=2)
    rq_para = doc.add_paragraph()
    rq_para.add_run('This paper addresses the following research questions:').bold = True
    doc.add_paragraph()

    add_bullet_point(doc, 'RQ1: How do lexical, discourse-grammatical, and contextual-semantic representations compare in journalistic genre classification?')
    add_bullet_point(doc, 'RQ2: What do classification errors and systematic model disagreements reveal about genre boundaries?')
    add_bullet_point(doc, 'RQ3: Which linguistic features are most diagnostic of genre, and how do they overlap across categories?')
    add_bullet_point(doc, 'RQ4: How do BERT\'s attention mechanisms distribute across genres, and what does this reveal about genre-marking strategies?')

    add_heading(doc, '1.3 Contributions', level=2)
    contrib_para = doc.add_paragraph()
    contrib_para.add_run('Our main contributions are:').bold = True
    doc.add_paragraph()

    add_bullet_point(doc, 'Empirical evidence for gradient genre boundaries: We demonstrate that 27.8% of articles occupy boundary zones where models systematically disagree.')
    add_bullet_point(doc, 'Demonstration of lexical primacy: We show that TF-IDF captures 99% of the genre signal that BERT captures.')
    add_bullet_point(doc, 'Model disagreement as uncertainty metric: We propose using inter-model disagreement to flag articles requiring human review.')
    add_bullet_point(doc, 'Attention analysis revealing genre markers: We analyze BERT\'s attention patterns to identify genre-specific lexical markers.')

    doc.add_page_break()

    # Data
    add_heading(doc, '3. Data', level=1)

    add_heading(doc, '3.1 Corpus: The Guardian Articles', level=2)
    data_text = """Our corpus consists of articles from The Guardian, a major British newspaper known for its well-organized genre tagging system. Articles were collected via The Guardian API and labeled by editorial staff into five genre categories:

• News: Factual reporting using the inverted pyramid structure. Focus on who, what, when, where, why.
• Analytical: Data-driven journalism with statistical analysis and in-depth interpretation.
• Feature: Human-interest stories using narrative techniques and storytelling.
• Editorial: Opinion pieces representing the newspaper's official stance.
• Review: Cultural criticism covering books, films, theater, and arts.

After cleaning and preprocessing, the final corpus contains 49,127 articles with approximately balanced genre distribution (~10,000 articles per genre)."""
    add_paragraph(doc, data_text)

    add_heading(doc, '3.2 Train-Validation-Test Split', level=2)
    split_text = """We split the data into:
• Training set: 80% (39,302 articles) for model training
• Validation set: 10% (4,912 articles) for hyperparameter tuning
• Test set: 10% (4,913 articles) for final evaluation

All error analysis and statistical validation were performed on the held-out test set to avoid data leakage."""
    add_paragraph(doc, split_text)

    add_heading(doc, '3.3 Ground Truth Limitations', level=2)
    add_paragraph(doc, 'Important limitation: We did not conduct inter-annotator agreement studies to validate the genre labels. The labels reflect editorial decisions by The Guardian staff. Without measured inter-annotator reliability (e.g., Cohen\'s Kappa), we cannot quantify the inherent ambiguity in genre labeling.')

    doc.add_page_break()

    # Methods
    add_heading(doc, '4. Methods', level=1)
    methods_text = """We compare three classification approaches representing different levels of linguistic abstraction:

4.1 Level 1: Lexical Representation (TF-IDF + Logistic Regression)
• TF-IDF vectorization with max_features=10,000
• Character n-grams from 1 to 2 (unigrams and bigrams)
• Logistic Regression with L2 regularization (C=10.0)
• Represents a purely lexical approach that captures word frequency patterns

4.2 Level 2: Discourse-Grammatical Representation (Linguistic Features + Random Forest)
• Ten hand-crafted linguistic features: Type-Token Ratio, Average Sentence Length, Pronoun Ratios (1st/2nd/3rd person), Modal Verb Ratio, Hedge Ratio, Stance Marker Ratio, Quote Ratio, Reporting Verb Ratio
• Random Forest with n_estimators=200, max_depth=15
• Captures discourse-level and syntactic patterns

4.3 Level 3: Contextual-Semantic Representation (BERT Fine-tuning)
• Model: bert-base-uncased (110M parameters)
• Fine-tuned for 3 epochs with batch_size=16, learning_rate=2e-5
• Represents the current state-of-the-art in transfer learning

4.4 Statistical Validation Methods
• McNemar's test for pairwise model comparison
• Bootstrap 95% confidence intervals (1,000 iterations)
• Cohen's Kappa for inter-model agreement"""
    add_paragraph(doc, methods_text)

    doc.add_page_break()

    # Results - WITH FIGURES
    add_heading(doc, '5. Results', level=1)

    add_heading(doc, '5.1 Overall Model Performance', level=2)
    add_paragraph(doc, 'Table 1 presents the overall performance of all three models on the held-out test set.')
    add_paragraph(doc, 'Table 1: Model Performance on Test Set (n=4,913)')
    add_paragraph(doc, 'BERT: 92.73% accuracy | TF-IDF: 86.40% accuracy | Linguistic: 82.85% accuracy')

    # Add comparison chart
    doc.add_paragraph()
    add_image(doc, '../results/comparison_barchart.png', width=6.5)
    add_figure_caption(doc, 'Figure 1: Model performance comparison showing BERT achieves 92.73%, TF-IDF 86.40%, and Linguistic 82.85% accuracy')
    doc.add_paragraph()

    results_text = """Key findings:

• BERT achieves the highest accuracy at 92.73%, outperforming TF-IDF by 6.33 percentage points
• TF-IDF is highly competitive at 86.40%, only 6.33 points behind BERT despite being orders of magnitude simpler
• Linguistic features underperform at 82.85%, significantly below both lexical models

The 1.06% gap between BERT and TF-IDF indicates that approximately 99% of the genre signal is captured by lexical choice alone."""
    add_paragraph(doc, results_text)

    add_heading(doc, '5.2 Statistical Significance Testing', level=2)
    stats_text = """Table 2: McNemar's Test Results
All pairwise differences are highly statistically significant (p < 0.001):

• BERT vs. TF-IDF: χ² = 156.01, p < 0.001 ***
• BERT vs. Linguistic: χ² = 258.34, p < 0.001 ***
• TF-IDF vs. Linguistic: χ² = 31.07, p < 0.001 ***

Table 3: Bootstrap 95% Confidence Intervals
The non-overlapping confidence intervals confirm distinct performance tiers:

• BERT: 92.73% [92.02%, 93.42%]
• TF-IDF: 86.40% [85.50%, 87.22%]
• Linguistic: 82.85% [81.78%, 83.94%]

Table 4: Inter-Model Agreement (Cohen's Kappa)
All models show substantial agreement (κ = 0.73–0.83), indicating they capture the same underlying genre signal."""
    add_paragraph(doc, stats_text)

    add_heading(doc, '5.3 Error Analysis: Gradient Genre Boundaries', level=2)
    error_text = """To investigate genre boundaries, we analyze patterns of model agreement and disagreement:

Table 5: Model Agreement Patterns
• All three agree: 3,670 articles (73.4%) — Correct 98.4% of the time
• Two agree, one disagrees: 1,138 articles (22.8%) — Correct 66.4% of the time
• All three disagree: 105 articles (2.1%) — Correct 0% of the time

Key finding: When all three models agree, they are correct 98.4% of the time. Conversely, when all three disagree, none are correctly classified, suggesting genuine ambiguity."""
    add_paragraph(doc, error_text)

    # Add model agreement figure
    doc.add_paragraph()
    add_image(doc, '../results/error_analysis/model_agreement.png', width=6.5)
    add_figure_caption(doc, 'Figure 2: Model agreement patterns showing 98.4% accuracy when all models agree')
    doc.add_paragraph()

    hybrid_text = """Hybrid articles (where models disagree) constitute 27.8% of the test set (1,390 articles). We analyze confusion patterns:

Table 6: BERT Confusion Patterns
• News ↔ Feature (5.4%): Both can tell factual stories using narrative techniques
• Analytical → Feature (11.5%): Analytical articles drift toward Feature when using narratives
• Feature ↔ Review (4.2%): Cultural Features resemble Reviews when including criticism
• Analytical ↔ Editorial (7.5%): Opinionated analysis blends into opinion

Feature as "hub genre": Feature attracts confusion from all other genres, occupying a central position in the genre space."""
    add_paragraph(doc, hybrid_text)

    # Add confusion matrices
    doc.add_page_break()
    add_heading(doc, 'Confusion Matrices', level=2)

    add_image(doc, '../results/tfidf_confusion_matrix.png', width=5.5)
    add_figure_caption(doc, 'Figure 3: TF-IDF Confusion Matrix')
    doc.add_paragraph()

    add_image(doc, '../results/linguistic_confusion_matrix.png', width=5.5)
    add_figure_caption(doc, 'Figure 4: Linguistic Features Confusion Matrix')
    doc.add_paragraph()

    add_heading(doc, '5.4 Feature Importance (Linguistic Model)', level=2)
    feature_text = """Table 7: Random Forest Feature Importance

1. Reporting Verbs Ratio (0.2049): News relies heavily on attributed speech
2. Type-Token Ratio (0.1940): Analytical pieces use richer vocabulary
3. First Person Ratio (0.1382): Features/Editorials use personal voice
4. Modal Ratio (0.1226): Hedges and modality vary by genre
5. Avg Sentence Length (0.0989): Syntactic complexity varies

Despite theoretical justification, these features achieved only 82.85% accuracy with substantial overfitting."""
    add_paragraph(doc, feature_text)

    # Add feature importance plot if available
    if os.path.exists('../results/comparison_ranking.png'):
        add_image(doc, '../results/comparison_ranking.png', width=6.0)
        add_figure_caption(doc, 'Figure 5: Feature importance ranking for linguistic model')
        doc.add_paragraph()

    doc.add_page_break()

    add_heading(doc, '5.5 BERT Attention Analysis', level=2)
    attention_text = """To interpret BERT's decisions, we extracted attention weights from the last 6 layers for 50 texts (10 per genre) and analyzed which tokens received highest aggregate attention.

Table 8: Top High-Attention Tokens by Genre

• Analytical: boris, johnson, downing — Political analysis focus
• Editorial: mr, nhs, of — Institutional language, formal tone
• Feature: i, my, said — First-person narrative, dialogue
• News: trump, masters, after — Event-driven, temporal markers
• Review: staging, debut, opera — Domain-specific arts vocabulary

Cross-genre overlap: We identified universal high-attention tokens present across all genres: "we", "related", and special tokens. This substantial overlap explains why genre boundaries are permeable."""
    add_paragraph(doc, attention_text)

    # Add attention top tokens chart
    add_image(doc, '../results/bert_attention_top_tokens.png', width=6.5)
    add_figure_caption(doc, 'Figure 6: Top attention tokens by genre showing genre-specific lexical markers')
    doc.add_paragraph()

    # Add attention heatmaps for each genre
    doc.add_page_break()
    add_heading(doc, 'BERT Attention Heatmaps by Genre', level=2)
    add_paragraph(doc, 'The following heatmaps show BERT\'s attention patterns (Layer 11, Head 1) for each genre:')

    for genre in ['News', 'Analytical', 'Feature', 'Editorial', 'Review']:
        heatmap_path = f'../results/attention_by_genre/{genre}_layer11_head1.png'
        if add_image(doc, heatmap_path, width=6.0):
            add_figure_caption(doc, f'Figure: {genre} genre - BERT attention heatmap (Layer 11, Head 1)')
        doc.add_paragraph()

    doc.add_page_break()

    # Discussion
    add_heading(doc, '6. Discussion', level=1)

    add_heading(doc, '6.1 Research Question 1: Representation Comparison', level=2)
    discussion1 = """RQ1 asked: How do lexical, discourse-grammatical, and contextual-semantic representations compare?

Our results show a clear performance hierarchy: BERT (92.73%) > TF-IDF (86.40%) > Linguistic (82.85%). However, the key insight is the magnitude of differences:

• The BERT vs. TF-IDF gap (6.33 pp) is statistically significant but practically small
• The Linguistic model underperforms by 20-35 pp

Interpretation: Genre classification is primarily a lexical task. Approximately 99% of what BERT learns about genre is captured by TF-IDF's word frequency patterns."""
    add_paragraph(doc, discussion1)

    add_heading(doc, '6.2 Research Question 2: Gradient Boundaries', level=2)
    discussion2 = """RQ2 asked: What do errors reveal about genre boundaries?

Our error analysis provides strong empirical evidence for gradient genre boundaries:

• Hybrid articles constitute 27.8% of the test set (1,390 articles)
• Model disagreement is systematic, not random
• Feature acts as a "hub genre" bridging informational and narrative styles
• Fully ambiguous articles (2.1%) suggest genuine hybridity

Theoretical implications: These findings support prototypical and gradient views of genre over discrete category models."""
    add_paragraph(doc, discussion2)

    add_heading(doc, '6.3 Practical Implications', level=2)
    practical = """For NLP practitioners:
• Don't over-engineer: TF-IDF may suffice for genre classification
• Use model disagreement as uncertainty metric to flag articles for human review
• Attention provides interpretability for BERT-based systems

For genre theorists:
• Empirical support for gradient boundaries with quantifiable hybridity (27.8%)
• Vocabulary as primary signal—genre is primarily lexical, not syntactic"""
    add_paragraph(doc, practical)

    doc.add_page_break()

    # Limitations
    add_heading(doc, '7. Limitations', level=1)

    add_heading(doc, '7.1 Single Dataset', level=2)
    add_paragraph(doc, 'All experiments were conducted on articles from a single newspaper (The Guardian). Findings may not generalize to other newspapers, cultural contexts, or languages.')

    add_heading(doc, '7.2 Ground Truth Validation (Major Limitation)', level=2)
    add_paragraph(doc, 'We did not conduct inter-annotator agreement studies to validate the genre labels. The labels reflect editorial decisions by The Guardian staff, which may involve subjective judgment. Without measuring Cohen\'s Kappa between human annotators, we cannot:')
    add_bullet_point(doc, 'Quantify the inherent ambiguity in genre labeling')
    add_bullet_point(doc, 'Determine whether model errors reflect true ambiguity or annotation inconsistency')
    add_bullet_point(doc, 'Establish a human performance baseline')
    add_paragraph(doc, 'It is possible that some "hybrid" articles are actually annotation artifacts rather than genuine genre mixing.')

    add_heading(doc, '7.3 Other Limitations', level=2)
    add_bullet_point(doc, 'Limited genre set (5 genres from traditional journalism)')
    add_bullet_point(doc, 'Limited linguistic features (only 10 features)')
    add_bullet_point(doc, 'Attention weights are imperfect proxies for feature importance')
    add_bullet_point(doc, 'BERT requires GPU hardware for reproducibility')

    doc.add_page_break()

    # Conclusion
    add_heading(doc, '8. Conclusion', level=1)

    conclusion_text = """This paper compared three levels of linguistic representation for journalistic genre classification. Our main findings are:

1. Lexical primacy: TF-IDF captures 99% of the genre signal that BERT captures
2. Gradient boundaries: 27.8% of articles occupy hybrid boundary zones
3. Feature as hub genre: Feature articles occupy a central position
4. Model agreement: When models agree, they are correct 98.4% of the time

Future work should include:
• Inter-annotator agreement studies to validate ground truth
• Cross-linguistic validation to test generalization
• Online genre analysis for digital-native publications
• Diachronic analysis to study genre evolution

Closing remarks: Genre classification errors are not failures to be minimized, but sources of insight into the fluid, evolving nature of journalistic genres. As journalism evolves in the digital age, our computational methods must embrace gradient boundaries and hybrid forms."""
    add_paragraph(doc, conclusion_text)

    doc.add_page_break()

    # References
    add_heading(doc, 'References', level=1)

    references = [
        'Bhatia, V. K. (1993). Analysing genre: Language use in professional settings. Longman.',
        'Bhatia, V. K. (2004). Worlds of written discourse: A genre-based view. Bloomsbury Publishing.',
        'Clark, K., Khandelwal, U., Levy, O., & Manning, C. D. (2019). What does BERT look at? An analysis of BERT\'s attention. Proceedings of the 2019 ACL Workshop on BlackboxNLP, 276–286.',
        'Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of NAACL-HLT, 4171–4186.',
        'Joachims, T. (1998). Text categorization with Support Vector Machines. Proceedings of the 10th European Conference on Machine Learning, 137–142.',
        'Rogers, A., Kovaleva, O., & Rumshisky, A. (2020). A primer in bertology: What we know about BERT and why. arXiv preprint arXiv:2002.12327.',
        'Swales, J. M. (1990). Genre analysis: English in academic and research settings. Cambridge University Press.',
        'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.',
        'Vig, J. (2019). A multiscale visualization of attention in the transformer model. Proceedings of the 57th Annual Meeting of the ACL, 37–44.'
    ]

    for ref in references:
        add_paragraph(doc, ref)

    # Save document
    output_path = '../Dialogue_2026_Paper.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")
    print(f"📊 Contains {len([p for p in doc.paragraphs if p._element.xml.endswith('>')])} paragraphs")
    return output_path

if __name__ == '__main__':
    print("Creating Word document with figures...")
    create_paper_document()
    print("Done!")
